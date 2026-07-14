"""
Text Extractor

Extracts text content from various document formats:
- PDF (using pypdf)
- Plain text (UTF-8)
- DOCX (using python-docx)
"""

import io
from typing import List, Tuple

import structlog

logger = structlog.get_logger(__name__)


class ExtractionResult:
    """Result of text extraction from a document."""

    def __init__(
        self,
        text: str,
        pages: int = 1,
        metadata: dict = None,
    ):
        self.text = text
        self.pages = pages
        self.metadata = metadata or {}
        self.char_count = len(text)
        self.word_count = len(text.split())


class TextExtractor:
    """
    Extracts text from supported document formats.

    Supported formats:
    - application/pdf → pypdf
    - text/plain → direct UTF-8 decode
    - application/vnd.openxmlformats-officedocument.wordprocessingml.document → python-docx
    """

    SUPPORTED_TYPES = {
        "application/pdf": "_extract_pdf",
        "text/plain": "_extract_text",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "_extract_docx",
    }

    def extract(self, content: bytes, content_type: str, filename: str = "") -> ExtractionResult:
        """
        Extract text from document content.

        Args:
            content: Raw file bytes
            content_type: MIME type
            filename: Original filename (for metadata)

        Returns:
            ExtractionResult with extracted text and metadata

        Raises:
            ValueError: If content type is not supported
        """
        method_name = self.SUPPORTED_TYPES.get(content_type)
        if not method_name:
            raise ValueError(f"Unsupported content type: {content_type}")

        method = getattr(self, method_name)
        result = method(content)
        result.metadata["filename"] = filename
        result.metadata["content_type"] = content_type
        result.metadata["original_size_bytes"] = len(content)

        logger.info(
            "Text extracted",
            filename=filename,
            pages=result.pages,
            chars=result.char_count,
            words=result.word_count,
        )

        return result

    def _extract_pdf(self, content: bytes) -> ExtractionResult:
        """Extract text from PDF."""
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        pages = len(reader.pages)

        text_parts: List[str] = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        full_text = "\n\n".join(text_parts)

        return ExtractionResult(
            text=full_text,
            pages=pages,
            metadata={"format": "pdf"},
        )

    def _extract_text(self, content: bytes) -> ExtractionResult:
        """Extract text from plain text file."""
        # Try UTF-8 first, fall back to latin-1
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")

        return ExtractionResult(
            text=text,
            pages=1,
            metadata={"format": "text", "encoding": "utf-8"},
        )

    def _extract_docx(self, content: bytes) -> ExtractionResult:
        """Extract text from DOCX."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            # Fallback: try basic extraction without python-docx
            logger.warning("python-docx not installed, attempting basic extraction")
            return ExtractionResult(
                text="[DOCX extraction requires python-docx package]",
                pages=1,
                metadata={"format": "docx", "error": "python-docx not installed"},
            )

        doc = DocxDocument(io.BytesIO(content))
        paragraphs: List[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n\n".join(paragraphs)

        return ExtractionResult(
            text=full_text,
            pages=1,  # DOCX doesn't have page concept in extraction
            metadata={"format": "docx", "paragraphs": len(paragraphs)},
        )
